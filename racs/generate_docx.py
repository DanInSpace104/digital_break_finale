from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

# from .models import Claim


def generate_docx(claim):
    doc = Document()
    par = doc.add_paragraph(
        '''Заместителю генерального директора –
        Директору по развитию
        В.О. Акуличеву'''
    )
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    par = doc.add_heading('Заявление на рационализаторское предложение', level=1)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par = doc.add_paragraph(
        f'''Зарегистрировано
        № {claim.id}
    {claim.creation_dt}'''
    )
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    par = doc.add_paragraph(claim.users.all().first().profile.otdel.name)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    par = doc.add_table(rows=1, cols=2)
    par.rows[0].cells[0].text = 'Имя Сотрудника'
    par.rows[0].cells[1].text = 'Фамилия '
    for user in claim.users.all():
        cells = par.add_row().cells
        cells[0].text = user.first_name
        cells[1].text = user.last_name

    par = doc.add_paragraph(
        f'''
        Предлагается рассмотреть предложение под наименованием:
        {claim.name}
        признать его рационализаторским, принять к использованию и выдать удостоверение на рационализаторское предложение.

Категория предложения в части цифровой трансформации: {claim.category.all().first()}
    '''
    )

    par = doc.add_heading('ОПИСАНИЕ ПРЕДЛОЖЕНИЯ', level=1)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par = doc.add_paragraph(claim.curr_desc)

    par = doc.add_heading('ОПИСАНИЕ ПРЕДЛАГАЕМОГО РЕШЕНИЯ', level=1)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par = doc.add_paragraph(claim.new_desc)

    par = doc.add_heading(
        'OЖИДАЕМЫЙ ПОЛОЖИТЕЛЬНЫЙ ЭФФЕКТ ОТ ИСПОЛЬЗОВАНИЯ (технический, организационный, управленческий или иной)',
        level=1,
    )
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par = doc.add_paragraph(claim.pos_effect)

    par = doc.add_paragraph('Необходимые затраты на внедрение')
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'Наименование'
    table.rows[0].cells[1].text = 'Стоимость'
    for cost in claim.costs.all():
        cells = table.add_row().cells
        cells[0].text = cost.name
        cells[1].text = str(cost.summ)

    par = doc.add_paragraph('Требуемые сроки на внедрение')
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'Этап'
    table.rows[0].cells[1].text = 'Длительность'
    for term in claim.terms.all():
        cells = table.add_row().cells
        cells[0].text = term.name
        cells[1].text = str(term.days)

    par = doc.add_paragraph(
        f'''Настоящим подтверждается действительное авторство и в соответствии с творческим участием каждого из авторов заключается следующее соглашение:

        СОГЛАШЕНИЕ
        о распределении вознаграждения ( % ) за использование рационализаторского предложения
        № {claim.id} от {claim.creation_dt}
        '''
    )
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = 'ФИО'
    table.rows[0].cells[1].text = '% вознаграждения'
    table.rows[0].cells[2].text = 'Подпись'
    table.rows[0].cells[3].text = 'Дата'
    for user in claim.users.all():
        cells = table.add_row().cells
        cells[0].text = user.first_name + ' ' + user.last_name
        cells[1].text = str(100 / claim.users.count()) + '%'
        cells[2].text = ''
        cells[3].text = str(claim.creation_dt)

    doc.save('tmp.docx')


if __name__ == '__main__':
    generate_docx(1)
